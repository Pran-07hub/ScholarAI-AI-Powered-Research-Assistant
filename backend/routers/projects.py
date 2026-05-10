from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from typing import List, Optional
from pydantic import BaseModel
from models import Project, User, Paper
from auth import get_current_user
from beanie import PydanticObjectId, Link
import asyncio
from datetime import datetime
from utils.project_access import require_project_access, ProjectRole
from utils.sanitize import sanitize_plain_text
import io
import json
import re
import logging

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/projects", tags=["Projects"])

@router.post("/", response_model=Project)
async def create_project(project: Project, current_user: User = Depends(get_current_user)):
    project.user_id = current_user
    await project.insert()
    return project

@router.get("/", response_model=List[Project])
async def list_projects(current_user: User = Depends(get_current_user)):
    uid = current_user.id
    projects = await Project.find({
        "$or": [
            {"user_id.$id": uid},
            {"members.$id": uid},
        ]
    }).sort(-Project.updated_at).to_list()
    return projects

@router.get("/{project_id}", response_model=Project)
async def get_project(project_id: PydanticObjectId, current_user: User = Depends(get_current_user)):
    project = await Project.get(project_id)
    require_project_access(project, current_user)
    return project

@router.put("/{project_id}", response_model=Project)
async def update_project(project_id: PydanticObjectId, updated_project: Project, current_user: User = Depends(get_current_user)):
    project = await Project.get(project_id)
    require_project_access(project, current_user, min_role=ProjectRole.OWNER)

    project.name = sanitize_plain_text(updated_project.name)
    project.description = sanitize_plain_text(updated_project.description) if updated_project.description else None
    project.updated_at = datetime.utcnow()
    await project.save()
    return project

@router.delete("/{project_id}")
async def delete_project(project_id: PydanticObjectId, current_user: User = Depends(get_current_user)):
    project = await Project.get(project_id)
    require_project_access(project, current_user, min_role=ProjectRole.OWNER)
    await project.delete()

    try:
        import asyncio
        from services.vector_store import delete_project_collection
        await asyncio.get_event_loop().run_in_executor(None, delete_project_collection, str(project_id))
    except Exception as e:
        logger.warning("Could not delete ChromaDB collection for project %s: %s", project_id, e)

    return {"message": "Project deleted"}

class PaperCreate(BaseModel):
    title: str
    authors: List[str] = []
    abstract: Optional[str] = None
    publication_date: Optional[str] = None
    pdf_url: Optional[str] = None
    source: str = "manual"
    doi: Optional[str] = None
    
class BulkPaperRequest(BaseModel):
    papers: List[PaperCreate]


class PaperImportRequest(BaseModel):
    identifier: str  # DOI (10.xxxx/...) or arXiv ID (2301.12345 or arxiv:2301.12345)
    project_id: Optional[str] = None  # overridden by path param


@router.post("/{project_id}/papers/import")
async def import_paper(
    project_id: PydanticObjectId,
    request: PaperImportRequest,
    current_user: User = Depends(get_current_user),
):
    """
    Import a paper by DOI or arXiv ID.
    Fetches metadata from Crossref (DOI) or arXiv API and saves as a Paper.
    """
    import aiohttp

    project = await Project.get(project_id)
    require_project_access(project, current_user)

    identifier = request.identifier.strip()
    metadata: dict = {}

    # Detect identifier type
    arxiv_id = None
    doi = None

    if identifier.lower().startswith("arxiv:"):
        arxiv_id = identifier[6:].strip()
    elif re.match(r"^\d{4}\.\d{4,5}(v\d+)?$", identifier):
        arxiv_id = identifier
    elif identifier.lower().startswith("10."):
        doi = identifier
    else:
        raise HTTPException(status_code=400, detail="Provide a valid DOI (10.xxx/...) or arXiv ID (2301.12345)")

    async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=15)) as session:
        if arxiv_id:
            url = f"http://export.arxiv.org/api/query?id_list={arxiv_id}&max_results=1"
            try:
                async with session.get(url) as resp:
                    text = await resp.text()
                # Parse Atom XML
                import xml.etree.ElementTree as ET
                ns = {"atom": "http://www.w3.org/2005/Atom", "arxiv": "http://arxiv.org/schemas/atom"}
                root = ET.fromstring(text)
                entry = root.find("atom:entry", ns)
                if entry is None:
                    raise HTTPException(status_code=404, detail="arXiv paper not found")

                title = (entry.findtext("atom:title", "", ns) or "").strip().replace("\n", " ")
                abstract = (entry.findtext("atom:summary", "", ns) or "").strip()
                published = (entry.findtext("atom:published", "", ns) or "")[:10]
                authors = [a.findtext("atom:name", "", ns) for a in entry.findall("atom:author", ns)]
                link_el = entry.find("atom:link[@rel='alternate']", ns)
                pdf_url = ""
                for link in entry.findall("atom:link", ns):
                    if link.get("title") == "pdf":
                        pdf_url = link.get("href", "")
                        break

                metadata = {
                    "title": title,
                    "authors": authors,
                    "abstract": abstract,
                    "year": published[:4] if published else None,
                    "pdf_url": pdf_url,
                    "source": "arxiv",
                    "venue": "arXiv",
                }
            except HTTPException:
                raise
            except Exception as e:
                raise HTTPException(status_code=502, detail=f"Failed to fetch from arXiv: {e}")

        elif doi:
            crossref_url = f"https://api.crossref.org/works/{doi}"
            try:
                async with session.get(crossref_url, headers={"User-Agent": "ScholarAI/1.0"}) as resp:
                    if resp.status == 404:
                        raise HTTPException(status_code=404, detail="DOI not found in Crossref")
                    data = await resp.json()

                work = data.get("message", {})
                title_list = work.get("title", [])
                title = title_list[0] if title_list else "Unknown"
                abstract_raw = work.get("abstract", "")
                # Strip JATS XML tags from abstract
                abstract = re.sub(r"<[^>]+>", "", abstract_raw).strip()

                author_list = work.get("author", [])
                authors = [
                    f"{a.get('given', '')} {a.get('family', '')}".strip()
                    for a in author_list
                ]

                year = None
                pub_date = work.get("published-print") or work.get("published-online") or {}
                dp = pub_date.get("date-parts", [[]])[0]
                if dp:
                    year = str(dp[0])

                container = work.get("container-title", [])
                venue = container[0] if container else None

                metadata = {
                    "title": title,
                    "authors": authors,
                    "abstract": abstract,
                    "year": year,
                    "pdf_url": None,
                    "source": "doi",
                    "venue": venue,
                    "doi": doi,
                }
            except HTTPException:
                raise
            except Exception as e:
                raise HTTPException(status_code=502, detail=f"Failed to fetch from Crossref: {e}")

    # Parse year
    pub_date_obj = None
    year_str = str(metadata.get("year") or "")
    if year_str.isdigit() and len(year_str) == 4:
        try:
            pub_date_obj = datetime.fromisoformat(f"{year_str}-01-01")
        except Exception:
            pass

    paper = Paper(
        project_id=project,
        title=metadata.get("title") or "Imported Paper",
        authors=metadata.get("authors") or [],
        abstract=metadata.get("abstract"),
        publication_date=pub_date_obj,
        venue=metadata.get("venue"),
        pdf_url=metadata.get("pdf_url"),
        source=metadata.get("source", "import"),
    )
    await paper.insert()

    # Index in ChromaDB
    try:
        import asyncio as _asyncio
        from services.vector_store import index_papers
        await _asyncio.get_event_loop().run_in_executor(None, index_papers, str(project_id), [{
            "id": str(paper.id),
            "text": paper.abstract or paper.title,
            "title": paper.title,
            "authors": paper.authors,
            "source": paper.source,
            "pdf_url": paper.pdf_url or "",
        }])
    except Exception as e:
        logger.warning("ChromaDB indexing failed for imported paper %s: %s", paper.id, e)

    return paper


@router.post("/{project_id}/papers")
async def add_paper_to_project(project_id: PydanticObjectId, paper_data: PaperCreate, current_user: User = Depends(get_current_user)):
    project = await Project.get(project_id)
    require_project_access(project, current_user)
    
    logger.info("Received paper save request: title=%r source=%r", paper_data.title, paper_data.source)
    pub_date = None
    if paper_data.publication_date:
        try:
            pub_date = datetime.fromisoformat(paper_data.publication_date.replace("Z", "+00:00"))
        except Exception:
            pass

    paper = Paper(
        project_id=project,
        title=paper_data.title,
        authors=paper_data.authors,
        abstract=paper_data.abstract,
        publication_date=pub_date,
        pdf_url=paper_data.pdf_url,
        source=paper_data.source,
        doi=paper_data.doi,
    )
    await paper.insert()

    try:
        import asyncio
        from services.vector_store import index_papers
        await asyncio.get_event_loop().run_in_executor(None, index_papers, str(project_id), [{
            "id":      str(paper.id),
            "text":    paper.abstract or paper.title,
            "title":   paper.title,
            "authors": paper.authors,
            "source":  paper.source,
            "pdf_url": paper.pdf_url or "",
        }])
    except Exception as e:
        logger.warning("ChromaDB indexing failed for paper %s: %s", paper.id, e)

    return paper

@router.post("/{project_id}/papers/bulk")
async def add_papers_bulk(project_id: PydanticObjectId, request: BulkPaperRequest, current_user: User = Depends(get_current_user)):
    logger.info(f"Bulk saving {len(request.papers)} papers payload received.")
    project = await Project.get(project_id)
    require_project_access(project, current_user)
    
    saved_papers = []
    for p_data in request.papers:
        pub_date = None
        if p_data.publication_date:
            try:
                pub_date = datetime.fromisoformat(p_data.publication_date.replace("Z", "+00:00"))
            except Exception:
                pass
        
        paper = Paper(
            project_id=project,
            title=p_data.title,
            authors=p_data.authors,
            abstract=p_data.abstract,
            publication_date=pub_date,
            pdf_url=p_data.pdf_url,
            source=p_data.source,
            doi=p_data.doi,
        )
        await paper.insert()
        saved_papers.append(paper)
    
    try:
        import asyncio
        from services.vector_store import index_papers
        await asyncio.get_event_loop().run_in_executor(None, index_papers, str(project_id), [
            {
                "id":      str(p.id),
                "text":    p.abstract or p.title,
                "title":   p.title,
                "authors": p.authors,
                "source":  p.source,
                "pdf_url": p.pdf_url or "",
            }
            for p in saved_papers
        ])
    except Exception as e:
        logger.warning("ChromaDB bulk indexing failed for project %s: %s", project_id, e)

    return {"message": "Successfully saved papers.", "count": len(saved_papers), "papers": saved_papers}

@router.get("/{project_id}/papers", response_model=List[Paper])
async def list_project_papers(project_id: PydanticObjectId, current_user: User = Depends(get_current_user)):
    project = await Project.get(project_id)
    require_project_access(project, current_user)
    
    # Beanie handles Link resolution, but for query we might need to be careful.
    # If project_id store as DBRef/Link, simple equality might not work if we pass ObjectId.
    # However, usually Beanie manages this. 
    # Let's try matching by ID directly if possible, or fetch all and filter (inefficient but safe for now)
    # Better: Paper.find(Paper.project_id.id == project.id)
    
    papers = await Paper.find(Paper.project_id.id == project.id).to_list()
    return papers

class GapAnalysisRequest(BaseModel):
    paper_ids: List[str]

from services.gap_analysis import analyze_research_gaps

@router.post("/{project_id}/analyze-gaps")
async def analyze_project_gaps(project_id: PydanticObjectId, request: GapAnalysisRequest, current_user: User = Depends(get_current_user)):
    project = await Project.get(project_id)
    require_project_access(project, current_user)
    
    result = await analyze_research_gaps(request.paper_ids)
    return result

class BulkDeletePapersRequest(BaseModel):
    paper_ids: List[str]

@router.delete("/{project_id}/papers/bulk")
async def delete_papers_bulk(project_id: PydanticObjectId, request: BulkDeletePapersRequest, current_user: User = Depends(get_current_user)):
    project = await Project.get(project_id)
    require_project_access(project, current_user, min_role=ProjectRole.OWNER)
    
    # Convert string IDs to PydanticObjectId
    try:
        obj_ids = [PydanticObjectId(pid) for pid in request.paper_ids]
    except Exception:
         raise HTTPException(status_code=400, detail="Invalid paper ID format")

    # Delete papers that belong to this project
    deleted = await Paper.find({"_id": {"$in": obj_ids}, "project_id.$id": project.id}).delete()
    
    return {"message": f"Successfully deleted {deleted.deleted_count} papers."}


# ── Data Extraction endpoints ─────────────────────────────────────────────────

from services.data_extraction import extract_paper_data

@router.post("/{project_id}/papers/extract-all")
async def extract_all_papers(project_id: PydanticObjectId, current_user: User = Depends(get_current_user)):
    """Run AI extraction on all papers in the project and persist the results."""
    project = await Project.get(project_id)
    require_project_access(project, current_user)

    papers = await Paper.find(Paper.project_id.id == project.id).to_list()
    if not papers:
        return {"message": "No papers found in project.", "papers": []}

    extracted = await extract_paper_data(papers)

    # Persist extracted_data into each paper document
    updated_papers = []
    for paper in papers:
        pid_str = str(paper.id)
        if pid_str in extracted:
            paper.extracted_data = extracted[pid_str]
            await paper.save()
        updated_papers.append(paper)

    return {"message": "Extraction complete.", "papers": updated_papers}


class ExtractionUpdateRequest(BaseModel):
    field: str   # e.g. "sample_size"
    value: str

@router.patch("/{project_id}/papers/{paper_id}/extraction")
async def update_paper_extraction(
    project_id: PydanticObjectId,
    paper_id: PydanticObjectId,
    request: ExtractionUpdateRequest,
    current_user: User = Depends(get_current_user)
):
    """Save a manually edited extraction field for a specific paper."""
    project = await Project.get(project_id)
    require_project_access(project, current_user)

    paper = await Paper.get(paper_id)
    if not paper:
        raise HTTPException(status_code=404, detail="Paper not found")

    if paper.extracted_data is None:
        paper.extracted_data = {}
    paper.extracted_data[request.field] = request.value
    await paper.save()
    return {"message": "Extraction updated.", "extracted_data": paper.extracted_data}


# ── PDF Upload endpoint ───────────────────────────────────────────────────────

import pypdf
from langchain_google_genai import ChatGoogleGenerativeAI
from utils.research_paper_summariser.config import get_google_api_key


@router.post("/{project_id}/papers/upload-pdf")
async def upload_pdf_paper(
    project_id: PydanticObjectId,
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    """
    Upload a PDF, extract its text, use Gemini to infer metadata (title, authors,
    abstract, year), save as a Paper in MongoDB, and index in the FAISS vector store
    so the chat agent can query it.
    """
    project = await Project.get(project_id)
    require_project_access(project, current_user)

    if not (file.filename or "").lower().endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported")

    content = await file.read()
    if len(content) > 50 * 1024 * 1024:  # 50 MB cap
        raise HTTPException(status_code=400, detail="PDF exceeds the 50 MB size limit")

    # Extract text from PDF
    try:
        pdf_reader = pypdf.PdfReader(io.BytesIO(content))
        full_text = ""
        for page in pdf_reader.pages[:30]:
            full_text += (page.extract_text() or "") + "\n"
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse PDF: {str(e)}")

    if not full_text.strip():
        # Scanned/image PDF — try Gemini multimodal OCR on the first few pages
        logger.info("PDF text extraction empty — attempting Gemini multimodal OCR for %s", file.filename)
        try:
            import base64
            from google import genai as _genai
            _gclient = _genai.Client(api_key=get_google_api_key())
            pdf_b64 = base64.b64encode(content).decode()
            ocr_response = _gclient.models.generate_content(
                model="gemini-2.5-flash",
                contents=[
                    {
                        "parts": [
                            {"inline_data": {"mime_type": "application/pdf", "data": pdf_b64}},
                            {"text": (
                                "This is a scanned academic paper. "
                                "Please extract and return ALL readable text from it, including the title, "
                                "authors, abstract, and body text. Do not summarise — return the raw text."
                            )},
                        ]
                    }
                ],
            )
            full_text = ocr_response.text or ""
        except Exception as ocr_err:
            logger.warning("Gemini multimodal OCR failed: %s", ocr_err)

    if not full_text.strip():
        raise HTTPException(
            status_code=400,
            detail="Could not extract text from this PDF. It may be image-based or encrypted."
        )

    # Use Gemini to extract structured metadata
    llm = ChatGoogleGenerativeAI(
        model="gemini-2.5-flash",
        google_api_key=get_google_api_key(),
        temperature=0.1,
    )

    extraction_prompt = (
        "Extract metadata from this academic paper. "
        "Return ONLY a valid JSON object with exactly these fields:\n"
        '  "title": string\n'
        '  "authors": array of strings\n'
        '  "abstract": string (use the paper\'s own abstract; if absent, summarise the introduction in ≤150 words)\n'
        '  "year": string (4-digit year or "Unknown")\n\n'
        f"Paper text (opening portion):\n{full_text[:4000]}\n\nJSON:"
    )

    metadata: dict = {}
    try:
        response = await llm.ainvoke(extraction_prompt)
        response_text = response.content if hasattr(response, "content") else str(response)
        # Strip markdown code fences if present
        response_text = re.sub(r"```(?:json)?\s*|\s*```", "", response_text).strip()
        metadata = json.loads(response_text)
    except Exception as e:
        logger.warning(f"Gemini metadata extraction failed ({e}); falling back to filename")
        fallback_title = (file.filename or "Uploaded Paper").replace(".pdf", "").replace("_", " ").replace("-", " ")
        metadata = {
            "title": fallback_title,
            "authors": [],
            "abstract": full_text[:500],
            "year": "Unknown",
        }

    # Parse year into a datetime
    pub_date = None
    year_str = str(metadata.get("year", "") or "")
    if year_str.isdigit() and len(year_str) == 4:
        try:
            pub_date = datetime.fromisoformat(f"{year_str}-01-01")
        except Exception:
            pass

    paper = Paper(
        project_id=project,
        title=metadata.get("title") or (file.filename or "Uploaded Paper"),
        authors=metadata.get("authors") or [],
        abstract=metadata.get("abstract") or full_text[:500],
        publication_date=pub_date,
        pdf_url=None,
        source="pdf_upload",
    )
    await paper.insert()

    # Index the full extracted text in ChromaDB for per-project RAG
    try:
        import asyncio
        from services.vector_store import index_papers
        await asyncio.get_event_loop().run_in_executor(None, index_papers, str(project_id), [{
            "id":      str(paper.id),
            "text":    full_text[:8000],   # full text beats abstract for PDFs
            "title":   paper.title,
            "authors": paper.authors,
            "source":  "pdf_upload",
            "pdf_url": "",
        }])
    except Exception as e:
        logger.warning(f"Could not index PDF in ChromaDB: {e}")

    return paper


# ── Research Output Export endpoints ─────────────────────────────────────────

from fastapi.responses import PlainTextResponse, Response
import io


def _md_to_docx(title: str, md: str) -> bytes:
    """Convert a simple Markdown string to a .docx file and return bytes."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.core_properties.title = title

    for line in md.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            text = line[2:]
            # handle **bold** inline
            parts = text.split("**")
            for i, part in enumerate(parts):
                run = p.add_run(part)
                if i % 2 == 1:
                    run.bold = True
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            p.add_run(line[2:-2]).bold = True
        elif line.strip() == "":
            pass  # skip blank lines (headings already add spacing)
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@router.get("/{project_id}/export/gap-analysis")
async def export_gap_analysis(
    project_id: PydanticObjectId,
    format: str = "md",
    current_user: User = Depends(get_current_user),
):
    """Export the project's research gap analysis as Markdown or Word (.docx)."""
    project = await Project.get(project_id)
    require_project_access(project, current_user)

    papers = await Paper.find(Paper.project_id.id == project.id).to_list()
    if not papers:
        raise HTTPException(status_code=400, detail="No papers in project")

    from services.synthesis import analyze_research_gaps_from_papers
    result = await analyze_research_gaps_from_papers(papers)
    if not result:
        raise HTTPException(status_code=500, detail="Gap analysis failed")

    title = f"Research Gap Analysis — {project.name}"
    md = f"# {title}\n\n"
    md += f"## Current State of the Field\n\n{result.get('current_state', '')}\n\n"

    limitations = result.get("limitations", [])
    if limitations:
        md += "## Key Limitations\n\n" + "\n".join(f"- {l}" for l in limitations) + "\n\n"

    underexplored = result.get("underexplored", [])
    if underexplored:
        md += "## Underexplored Areas\n\n" + "\n".join(f"- {u}" for u in underexplored) + "\n\n"

    future = result.get("future_directions", [])
    if future:
        md += "## Future Research Directions\n\n" + "\n".join(f"- {f}" for f in future) + "\n\n"

    base = project.name.replace(" ", "_")[:40]
    if format == "docx":
        data = _md_to_docx(title, md)
        return Response(
            content=data,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{base}_gap_analysis.docx"'},
        )
    return PlainTextResponse(
        content=md,
        media_type="text/markdown",
        headers={"Content-Disposition": f'attachment; filename="{base}_gap_analysis.md"'},
    )


@router.get("/{project_id}/export/synthesis")
async def export_synthesis(
    project_id: PydanticObjectId,
    format: str = "md",
    current_user: User = Depends(get_current_user),
):
    """Export the project's topic synthesis as Markdown or Word (.docx)."""
    project = await Project.get(project_id)
    require_project_access(project, current_user)

    papers = await Paper.find(Paper.project_id.id == project.id).to_list()
    if not papers:
        raise HTTPException(status_code=400, detail="No papers in project")

    from services.synthesis import synthesize_topics
    topics = await synthesize_topics(papers)
    if not topics:
        raise HTTPException(status_code=500, detail="Synthesis failed")

    paper_map = {str(p.id): p for p in papers}

    title = f"Topic Synthesis — {project.name}"
    md = f"# {title}\n\n"
    for t in topics:
        md += f"## {t['topic']}\n\n{t['summary']}\n\n"
        paper_ids = t.get("paper_ids", [])
        if paper_ids:
            md += "**Papers:**\n"
            for pid in paper_ids:
                p = paper_map.get(pid)
                if p:
                    year = p.publication_date.year if p.publication_date else "n.d."
                    authors = ", ".join(p.authors[:2]) + (" et al." if len(p.authors) > 2 else "")
                    md += f"- {p.title} ({authors}, {year})\n"
            md += "\n"

    base = project.name.replace(" ", "_")[:40]
    if format == "docx":
        data = _md_to_docx(title, md)
        return Response(
            content=data,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{base}_synthesis.docx"'},
        )
    return PlainTextResponse(
        content=md,
        media_type="text/markdown",
        headers={"Content-Disposition": f'attachment; filename="{base}_synthesis.md"'},
    )


@router.get("/{project_id}/export/timeline")
async def export_timeline(
    project_id: PydanticObjectId,
    format: str = "md",
    current_user: User = Depends(get_current_user),
):
    """Export the project's research timeline as Markdown or Word (.docx)."""
    project = await Project.get(project_id)
    require_project_access(project, current_user)

    papers = await Paper.find(Paper.project_id.id == project.id).to_list()
    if not papers:
        raise HTTPException(status_code=400, detail="No papers in project")

    from services.timeline_service import generate_timeline
    timeline = await generate_timeline(papers, field_name=project.name)

    title = f"Research Timeline — {project.name}"
    md = f"# {title}\n\n"
    for entry in timeline:
        year = entry.get("year", "Unknown")
        summary = entry.get("summary", "")
        year_papers = entry.get("papers", [])
        md += f"## {year}\n\n{summary}\n\n"
        if year_papers:
            for p in year_papers:
                md += f"- **{p.get('title', '')}**"
                authors = p.get("authors", [])
                if authors:
                    md += f" — {', '.join(authors[:2])}"
                md += "\n"
            md += "\n"

    base = project.name.replace(" ", "_")[:40]
    if format == "docx":
        data = _md_to_docx(title, md)
        return Response(
            content=data,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{base}_timeline.docx"'},
        )
    return PlainTextResponse(
        content=md,
        media_type="text/markdown",
        headers={"Content-Disposition": f'attachment; filename="{base}_timeline.md"'},
    )


class ConvertRequest(BaseModel):
    title: str
    content_md: str
    filename: str = "export"


@router.post("/export/convert-docx")
async def convert_md_to_docx(
    body: ConvertRequest,
    current_user: User = Depends(get_current_user),
):
    """Convert a Markdown string to a .docx file and return it for download."""
    data = _md_to_docx(body.title, body.content_md)
    fname = body.filename.replace(" ", "_")[:60] + ".docx"
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{fname}"'},
    )


# ── Comparative Analysis endpoint ────────────────────────────────────────────

class CompareRequest(BaseModel):
    paper_ids: List[str]  # 2–5 paper IDs


@router.post("/{project_id}/compare-papers")
async def compare_papers(
    project_id: PydanticObjectId,
    request: CompareRequest,
    current_user: User = Depends(get_current_user),
):
    """
    Compare 2–5 papers across methodology, findings, sample size, and limitations.
    Returns a comparison matrix and an AI-generated summary.
    """
    from google import genai as _genai_new
    from utils.json_repair import parse_json_robust

    project = await Project.get(project_id)
    require_project_access(project, current_user)

    if len(request.paper_ids) < 2 or len(request.paper_ids) > 5:
        raise HTTPException(status_code=400, detail="Select between 2 and 5 papers to compare")

    # Load the requested papers, verify they belong to this project
    papers = []
    for pid_str in request.paper_ids:
        try:
            paper = await Paper.get(PydanticObjectId(pid_str))
        except Exception:
            paper = None
        if not paper:
            raise HTTPException(status_code=404, detail=f"Paper {pid_str} not found")
        paper_proj_id = str(paper.project_id.ref.id if isinstance(paper.project_id, Link) else paper.project_id.id)
        if paper_proj_id != str(project_id):
            raise HTTPException(status_code=403, detail=f"Paper {pid_str} does not belong to this project")
        papers.append(paper)

    # Build compact paper descriptions for the prompt
    paper_summaries = []
    for i, p in enumerate(papers):
        year = p.publication_date.year if p.publication_date else "n.d."
        authors_str = ", ".join(p.authors[:3]) + (" et al." if len(p.authors) > 3 else "")
        abstract = (p.abstract or "")[:600]
        paper_summaries.append(
            f"Paper {i + 1}: {p.title} ({authors_str}, {year})\nAbstract: {abstract}"
        )

    prompt = (
        "You are a research synthesis assistant. Compare the following papers and return a JSON object.\n\n"
        + "\n\n".join(paper_summaries)
        + "\n\nReturn ONLY a valid JSON object with this structure:\n"
        '{\n'
        '  "dimensions": ["Methodology", "Sample Size", "Key Findings", "Limitations", "Dataset/Context"],\n'
        '  "matrix": [\n'
        '    {"paper_index": 1, "title": "...", "values": {"Methodology": "...", "Sample Size": "...", "Key Findings": "...", "Limitations": "...", "Dataset/Context": "..."}},\n'
        '    ...\n'
        '  ],\n'
        '  "summary": "A 2-3 sentence overall comparison highlighting key similarities and differences."\n'
        '}'
    )

    api_key = get_google_api_key()
    _gclient_cmp = _genai_new.Client(api_key=api_key)

    try:
        response = await asyncio.get_event_loop().run_in_executor(
            None, lambda: _gclient_cmp.models.generate_content(
                model="gemini-2.5-flash", contents=prompt
            )
        )
        raw = response.text or ""
        raw = re.sub(r"```(?:json)?\s*|\s*```", "", raw).strip()
        result = parse_json_robust(raw)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Comparison generation failed: {e}")

    if not result or "matrix" not in result:
        raise HTTPException(status_code=500, detail="Could not parse comparison result")

    # Attach paper metadata to matrix rows
    paper_meta = {str(i + 1): {"id": str(p.id), "title": p.title, "authors": p.authors[:3]} for i, p in enumerate(papers)}
    for row in result.get("matrix", []):
        idx = str(row.get("paper_index", ""))
        if idx in paper_meta:
            row["id"] = paper_meta[idx]["id"]
            row["authors"] = paper_meta[idx]["authors"]

    return result


# ── Topic Synthesis endpoint ──────────────────────────────────────────────────

from services.synthesis import synthesize_topics

@router.post("/{project_id}/synthesize")
async def synthesize_project_topics(project_id: PydanticObjectId, current_user: User = Depends(get_current_user)):
    """Use AI to identify main topics across all project papers and map which papers cover each topic."""
    project = await Project.get(project_id)
    require_project_access(project, current_user)

    papers = await Paper.find(Paper.project_id.id == project.id).to_list()
    if not papers:
        return {"topics": [], "papers": []}

    topics = await synthesize_topics(papers)

    # Return topics + a flat list of paper metadata for the frontend to look up by id
    papers_meta = [
        {
            "id": str(p.id),
            "title": p.title,
            "authors": p.authors,
            "publication_date": p.publication_date.isoformat() if p.publication_date else None,
        }
        for p in papers
    ]

    return {"topics": topics, "papers": papers_meta}


# ── Research Timeline endpoint ────────────────────────────────────────────────

@router.get("/{project_id}/dashboard")
async def project_dashboard(
    project_id: PydanticObjectId,
    current_user: User = Depends(get_current_user),
):
    """Return a dashboard summary for a project: counts, members, recent activity."""
    from models import Note, PaperAnnotation
    from beanie import Link as BeanieLink

    project = await Project.get(project_id)
    require_project_access(project, current_user)

    # Counts
    papers = await Paper.find(Paper.project_id.id == project.id).to_list()
    notes = await Note.find(Note.project_id.id == project.id).sort(-Note.updated_at).to_list()
    annotations = await PaperAnnotation.find(
        PaperAnnotation.project_id.id == project.id
    ).sort(-PaperAnnotation.created_at).to_list()

    # Recent activity — last 5 papers + notes combined, sorted by created_at
    recent_items = []
    for p in papers[-5:]:
        recent_items.append({
            "type": "paper",
            "title": p.title,
            "authors": p.authors[:2],
            "created_at": p.created_at.isoformat(),
            "id": str(p.id),
        })
    for n in notes[:5]:
        recent_items.append({
            "type": "note",
            "title": n.title,
            "created_at": n.updated_at.isoformat(),
            "id": str(n.id),
        })
    recent_items.sort(key=lambda x: x["created_at"], reverse=True)

    # Members
    member_ids = []
    for m in project.members or []:
        mid = str(m.ref.id if isinstance(m, BeanieLink) else m.id)
        member_ids.append(mid)

    # Owner info
    owner_id = str(project.user_id.ref.id if isinstance(project.user_id, BeanieLink) else project.user_id.id)

    # Sources breakdown
    source_counts: dict = {}
    for p in papers:
        source_counts[p.source] = source_counts.get(p.source, 0) + 1

    return {
        "project": {
            "id": str(project.id),
            "name": project.name,
            "description": project.description,
            "created_at": project.created_at.isoformat(),
            "updated_at": project.updated_at.isoformat(),
            "owner_id": owner_id,
            "member_count": len(member_ids),
        },
        "stats": {
            "paper_count": len(papers),
            "note_count": len(notes),
            "annotation_count": len(annotations),
            "sources": source_counts,
        },
        "recent_activity": recent_items[:8],
    }


from services.timeline_service import generate_timeline

@router.get("/{project_id}/timeline")
async def research_timeline(project_id: PydanticObjectId, current_user: User = Depends(get_current_user)):
    """Build a year-by-year research timeline from project papers with AI synthesis."""
    project = await Project.get(project_id)
    require_project_access(project, current_user)

    papers = await Paper.find(Paper.project_id.id == project.id).to_list()
    if not papers:
        return {"timeline": []}

    timeline = await generate_timeline(papers, field_name=project.name)
    return {"timeline": timeline}


